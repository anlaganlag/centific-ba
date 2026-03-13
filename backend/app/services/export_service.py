import json
import io
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.analysis import Feature, UserStory, AcceptanceCriterion


class ExportService:
    """Generates DOCX exports from analysis results."""

    def generate_docx(self, features_json: str, project_name: str = "Project") -> io.BytesIO:
        """
        Generate a DOCX document from analysis features.

        Args:
            features_json: JSON string of Feature list
            project_name: Name of the project for the title

        Returns:
            BytesIO buffer containing the DOCX file
        """
        features = [Feature(**f) for f in json.loads(features_json)]

        doc = Document()

        # Title
        title = doc.add_heading(f"{project_name} — BA Analysis Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Features overview
        doc.add_heading("Features Overview", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Feature"
        hdr[2].text = "Problem"
        hdr[3].text = "Business Process"

        for feat in features:
            row = table.add_row().cells
            row[0].text = feat.feature_id
            row[1].text = feat.title
            row[2].text = feat.problem_statement[:80] + ("..." if len(feat.problem_statement) > 80 else "")
            row[3].text = feat.business_process

        doc.add_page_break()

        # Detailed features
        for feat in features:
            doc.add_heading(f"{feat.feature_id}: {feat.title}", level=1)

            # Feature details
            doc.add_heading("Problem Statement", level=2)
            doc.add_paragraph(feat.problem_statement)

            doc.add_heading("Benefit", level=2)
            doc.add_paragraph(feat.benefit)

            doc.add_heading("Business Process", level=2)
            doc.add_paragraph(feat.business_process)

            doc.add_heading("Scope", level=2)
            doc.add_paragraph(feat.scope)

            if feat.sources:
                doc.add_heading("Sources", level=2)
                for src in feat.sources:
                    doc.add_paragraph(src, style="List Bullet")

            # User stories
            if feat.user_stories:
                doc.add_heading("User Stories", level=2)

                for story in feat.user_stories:
                    # Story title
                    p = doc.add_paragraph()
                    run = p.add_run(f"{story.story_id}: ")
                    run.bold = True
                    p.add_run(
                        f"As a {story.as_a}, I want {story.i_want}, so that {story.so_that}"
                    )

                    # Acceptance criteria
                    if story.acceptance_criteria:
                        doc.add_paragraph("Acceptance Criteria:", style="List Bullet")
                        for ac in story.acceptance_criteria:
                            criterion_text = f"Given {ac.given}, When {ac.when}, Then {ac.then}"
                            doc.add_paragraph(criterion_text, style="List Bullet 2")

                    # Business rules
                    if story.business_rules:
                        doc.add_paragraph("Business Rules:", style="List Bullet")
                        for rule in story.business_rules:
                            doc.add_paragraph(rule, style="List Bullet 2")

                    # Dependencies
                    if story.dependencies:
                        doc.add_paragraph(
                            f"Dependencies: {', '.join(story.dependencies)}",
                            style="List Bullet"
                        )

                    doc.add_paragraph("")  # spacing

            doc.add_page_break()

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
